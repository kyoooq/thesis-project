import docx  # python-docx


def extract_text_from_docx(docx_path: str) -> str:

    document = docx.Document(docx_path)
    parts = []

    # Body paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cell text
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)
